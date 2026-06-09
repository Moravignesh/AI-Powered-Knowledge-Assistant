import os
import logging
from typing import Optional
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyPDF2, with OCR fallback."""
    text = ""
    try:
        import PyPDF2
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                text += f"\n[Page {page_num + 1}]\n{page_text}"

        # If extracted text is too short, try OCR
        if len(text.strip()) < 100:
            text = _ocr_pdf(file_path) or text

    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        try:
            text = _ocr_pdf(file_path) or ""
        except Exception as ocr_e:
            logger.error(f"OCR fallback error: {ocr_e}")

    return text.strip()


def _ocr_pdf(file_path: str) -> Optional[str]:
    """OCR a PDF using pytesseract and PIL."""
    try:
        import pytesseract
        from PIL import Image
        import fitz  # PyMuPDF - optional, only if installed

        text = ""
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            mat = fitz.Matrix(2, 2)  # 2x zoom for better OCR
            pix = page.get_pixmap(matrix=mat)
            img_path = f"/tmp/ocr_page_{page_num}.png"
            pix.save(img_path)
            img = Image.open(img_path)
            page_text = pytesseract.image_to_string(img)
            text += f"\n[Page {page_num + 1}]\n{page_text}"
            os.remove(img_path)
        return text.strip()
    except ImportError:
        # Try simpler OCR without fitz
        try:
            import pytesseract
            from PIL import Image
            from pdf2image import convert_from_path

            images = convert_from_path(file_path)
            text = ""
            for i, image in enumerate(images):
                text += f"\n[Page {i + 1}]\n{pytesseract.image_to_string(image)}"
            return text.strip()
        except Exception:
            return None
    except Exception as e:
        logger.error(f"OCR error: {e}")
        return None


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    try:
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        return ""
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        return ""


def extract_text(file_path: str, file_type: str) -> str:
    """Route to appropriate extractor based on file type."""
    extractors = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "txt": extract_text_from_txt,
    }
    extractor = extractors.get(file_type.lower())
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")
    return extractor(file_path)


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks for embedding."""
    if not text:
        return []

    chunks = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + chunk_size, text_len)

        # Try to end at sentence boundary
        if end < text_len:
            for sep in [". ", ".\n", "\n\n", "\n", " "]:
                boundary = text.rfind(sep, start + chunk_size // 2, end)
                if boundary != -1:
                    end = boundary + len(sep)
                    break

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - overlap if end - overlap > start else end

    return chunks
